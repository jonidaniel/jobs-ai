"""
Orchestrates the generation of cover letters.

CLASSES:
    GeneratorAgent

FUNCTIONS (in order of workflow):
    generate_letters     (public)
    _write_letters        (internal)
"""

import os
import logging
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from jobsai.config.paths import COVER_LETTER_PATH
from jobsai.config.prompts import (
    GENERATOR_SYSTEM_PROMPT as SYSTEM_PROMPT,
    GENERATOR_USER_PROMPT as USER_PROMPT,
)

from jobsai.utils.llms import call_llm
from jobsai.utils.normalization import normalize_text

logger = logging.getLogger(__name__)


class GeneratorAgent:
    """Generate cover letters.

    Args:
        timestamp (str): The backend-wide timestamp for consistent file naming.
    """

    def __init__(self, timestamp: str):
        self.timestamp = timestamp

    # ------------------------------
    # Public interface
    # ------------------------------
    def generate_letters(
        self,
        job_analysis: str,
        profile: str,
        style: str,
    ) -> Document:
        """Produce a tailored job-application message based on the candidate's skills and the job analysis.

        Args:
            profile (str): The candidate profile text.
            job_analysis (str): The job analysis that contains instructions for what kind of cover letter to write.
            letter_style (str): The intended style/tone of the cover letter.

        Returns:
            Document: The final cover letter.
        """

        # Build the system prompt
        tone_instructions = {
            "Professional": "Write in a clear, respectful, concise, professional tone. Use well-structured paragraphs. Avoid exaggerations.",
            "friendly": "Write in a warm, positive tone but keep it professional.",
            "confident": "Write with a confident, proactive tone without sounding arrogant.",
        }
        base_style = tone_instructions.get(style, tone_instructions["Professional"])
        system_prompt = SYSTEM_PROMPT.format(base_style=base_style)

        # Build the user prompt
        user_prompt = USER_PROMPT.format(profile=profile, job_analysis=job_analysis)

        # Write the cover letters
        cover_letters = self._write_letters(system_prompt, user_prompt)

        return cover_letters

    # ------------------------------
    # Internal function
    # ------------------------------
    def _write_letters(self, system_prompt: str, user_prompt: str) -> Document:
        """
        Write the cover letters with proper formatting.

        Creates a Word document with:
        1. Contact information (top-right)
        2. Date (top-right)
        3. Recipient placeholder (top-right)
        4. LLM-generated cover letter body
        5. Signature placeholder (bottom-right)

        The document follows standard business letter formatting conventions.

        Args:
            system_prompt (str): System prompt defining LLM's role as cover letter writer
            user_prompt (str): User prompt containing profile and job analysis

        Returns:
            Document: The complete cover letter as a python-docx Document object
        """

        cover_letter = Document()

        # Add contact information at the top right
        # Format: website, LinkedIn, GitHub (blank line) email, phone
        contact_paragraph = cover_letter.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        contact_paragraph.add_run("ADD WEBSITE\n")
        contact_paragraph.add_run("ADD LINKEDIN\n")
        contact_paragraph.add_run("ADD GITHUB\n\n")
        contact_paragraph.add_run("ADD EMAIL\n")
        contact_paragraph.add_run("ADD PHONE\n\n")

        # Convert timestamp (YYYYMMDD_HHMMSS) to human-readable date
        # Format: "Month Day, Year" (e.g., "November 30, 2025")
        parsed_timestamp = datetime.strptime(self.timestamp, "%Y%m%d_%H%M%S")
        pretty_date = parsed_timestamp.strftime("%B %d, %Y")
        # Add the date (aligned right)
        cover_letter.add_paragraph(f"{pretty_date}\n").alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )

        # Add recipient placeholder (aligned right)
        # User should manually fill these in before sending
        cover_letter.add_paragraph("ADD RECRUITER/HIRING TEAM").alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )
        cover_letter.add_paragraph("ADD HIRING COMPANY/GROUP\n\n").alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )

        # Generate cover letter body using LLM
        # The LLM writes the actual cover letter content based on prompts
        raw = call_llm(system_prompt, user_prompt)
        # Normalize text (clean whitespace, line breaks, etc.)
        normalized = normalize_text(raw)
        # Insert the body into the document
        cover_letter.add_paragraph(normalized)

        # Add signature section (aligned right)
        cover_letter.add_paragraph("Best regards,").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cover_letter.add_paragraph("ADD YOUR NAME").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Save the document to disk
        filename = f"{self.timestamp}_cover_letter.docx"
        cover_letter.save(os.path.join(COVER_LETTER_PATH, filename))

        return cover_letter
