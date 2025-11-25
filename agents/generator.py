# ---------- GENERATOR AGENT ----------

# generate_letters
# _build_system_prompt
# _build_user_prompt
# _write_letter

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict

from docx import Document

from utils.llms import call_llm
from utils.normalization import normalize_text

from config.schemas import SkillProfile

logger = logging.getLogger(__name__)


class GeneratorAgent:
    """
    Generate cover letters.
    """

    def __init__(self, letters_path: Path, timestamp: str):
        """
        Construct GeneratorAgent.
        """

        self.letters_path = letters_path
        self.timestamp = timestamp

    # ------------------------------
    # Public interface
    # ------------------------------
    def generate_letters(
        self,
        skill_profile: SkillProfile,
        job_report: str,
        letter_style: str,
        contact_info: Dict,
        # employer: Optional[str] = None,
        # job_title: Optional[str] = None,
        # ) -> str:
    ):
        """
        Produce a tailored job-application message based on
        the candidate's skills and the job report.

        Args:
            skill_profile: the skill profile
            job_report: the job report
            letter_style: the intended style/tone of the cover letter

        Returns:
            output: the generated text
        """

        system_prompt = self._build_system_prompt(letter_style)
        user_prompt = self._build_user_prompt(skill_profile, job_report)

        self._write_letter(contact_info)

        logger.info(" GENERATING APPLICATION TEXT...")

        raw = call_llm(system_prompt, user_prompt)

        output = normalize_text(raw)

        logger.info(" APPLICATION TEXT GENERATION COMPLETED\n")

        # return output

    # ------------------------------
    # Internal functions
    # ------------------------------

    def _build_system_prompt(self, style: str) -> str:
        """
        Build the system prompt.

        Args:
            style: the style or tone for the cover letter

        Returns:
            system_prompt: the system prompt
        """

        tone_instructions = {
            "professional": (
                "Write in a clear, respectful, concise, professional tone. "
                "Use well-structured paragraphs. Avoid exaggerations."
            ),
            "friendly": ("Write in a warm, positive tone but keep it professional."),
            "confident": (
                "Write with a confident, proactive tone without sounding arrogant."
            ),
        }

        base_style = tone_instructions.get(style, tone_instructions["professional"])

        system_prompt = f"""
        You are a professional cover letter writer.
        Your goal is to produce polished text suitable for real-world job applications.
        Follow this style:
        {base_style}
        """

        return system_prompt

    def _build_user_prompt(
        self,
        skill_profile: SkillProfile,
        job_report: str,
        # employer: Optional[str],
        # job_title: Optional[str],
    ) -> str:
        """
        Build the user prompt.

        Args:
            skill_profile: the skill profile
            job_report: the job report

        Returns:
            user_prompt: the user prompt
        """

        # employer_text = f"Employer: {employer}\n" if employer else ""
        employer_text = f"Employer: empty\n"
        # title_text = f"Target job title: {job_title}\n" if job_title else ""
        title_text = f"Target job title: empty\n"

        user_prompt = f"""
        Generate a tailored job-application message.

        Candidate Skill Profile (JSON):
        {skill_profile.model_dump_json(indent=2)}

        Job Match Analysis:
        {job_report}

        {employer_text}{title_text}

        Instructions:
        - Produce a compelling but concise job-application message.
        - Highlight the candidate's relevant skills based on the report.
        - If employer or job title are given, tailor the message to them.
        - Keep it truthful, specific, and readable.
        """

        return user_prompt

    def _write_letter(self, contact_info: Dict):
        doc = Document()

        p = doc.add_paragraph()
        p.add_run(f"{contact_info.get("website")}\n")
        p.add_run(f"{contact_info.get("linkedin")}\n")
        p.add_run(f"{contact_info.get("github")}\n\n")
        p.add_run(f"{contact_info.get("email")}\n")
        p.add_run(f"{contact_info.get("phone")}\n\n")

        # Turn the timestamp into a 'pretty date'
        dt = datetime.strptime(self.timestamp, "%Y%m%d_%H%M%S")
        pretty_date = dt.strftime("%B %d, %Y")

        doc.add_paragraph(f"{pretty_date}\n")

        doc.add_paragraph("Hiring Team")
        doc.add_paragraph("Vuono Group\n\n")

        # Body paragraphs
        body = []

        for paragraph in body:
            doc.add_paragraph(paragraph)

        filename = f"{self.timestamp}_cover_letter.docx"

        path = os.path.join(self.letters_path, filename)

        # Save the cover letter to /data/cover_letters/
        doc.save(path)
